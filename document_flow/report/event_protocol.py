from odoo import _, models
from docx.shared import Pt, Mm


class EventProtocol(models.AbstractModel):
    _name = 'report.document_flow.event_protocol'
    _description = 'document_flow.event_protocol'
    _inherit = "report.report_docx.abstract"

    def generate_docx_report(self, doc, data, objs):
        event = objs[0]
        if event:
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_after = 0
            paragraph = doc.add_paragraph(
                _('Protocol %s from %s') % (event.name, event.date_start.strftime('%d.%m.%Y')))
            run = paragraph.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            paragraph = doc.add_paragraph(_('Were present:'))
            paragraph.paragraph_format.space_before = Mm(5)
            run = paragraph.runs[0]
            run.font.italic = True
            for member in event.member_ids:
                doc.add_paragraph(member.name)
            paragraph = doc.add_paragraph(_('Agenda of the event:'))
            paragraph.paragraph_format.space_before = Mm(5)
            run = paragraph.runs[0]
            run.font.italic = True
            run.font.underline = True
            counter = 1
            for question in event.question_ids:
                doc.add_paragraph('%s. %s' % (counter, question.name.striptags())).paragraph_format.space_before = Mm(
                    3)
                counter += 1
            paragraph = doc.add_paragraph(_('Decisions:'))
            paragraph.paragraph_format.space_before = Mm(5)
            run = paragraph.runs[0]
            run.font.italic = True
            run.font.underline = True
            counter = 1
            for decision in event.decision_ids:
                doc.add_paragraph('%s. %s' % (counter, decision.name.striptags())).paragraph_format.space_before = Mm(
                    3)
                if decision.responsible_id:
                    doc.add_paragraph(_('Ответственный: %s') % decision.responsible_id.name)
                if decision.date_deadline:
                    doc.add_paragraph(_('Срок исполнения: %s') % decision.date_deadline)
                counter += 1
